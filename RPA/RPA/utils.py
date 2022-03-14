import io
import os
import re
import nltk
from pkg_resources import yield_lines
import spacy
import pandas as pd
import docx2txt
import pdfplumber
from . import constants as cs
from spacy.matcher import Matcher
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords 
from datetime import datetime
from dateutil import relativedelta
import docx
import pandas as pd
from docx.api import Document
import requests
import json
import random
from lxml import html


def extract_text_from_pdf(pdf_path):
    
    text = ''
    with pdfplumber.open(pdf_path) as pdf:
        for i in range(len(pdf.pages)):
            page = pdf.pages[i]
            text += page.extract_text() 
    return text

def extract_text_from_docx(doc_path):
    #doc = docx.Document(doc_path)
    #l = len(doc.paragraphs)

    #doc_text = ''

    #for i in range(l):
        #doc_text = doc_text + doc.paragraphs[i].text 
  

    #return doc_text

    
    temp = docx2txt.process(doc_path)

    #text = [line.replace('\t', ' ') for line in temp.split('\n') if line]
    return  temp #+ table_text #' '.join(text)
    

def extract_text_from_doc(doc_path):
    
    try:
        try:
            import textract
        except ImportError:
            return ' '
        text = textract.process(doc_path).decode('utf-8')
        return text
    except KeyError:
        return ' '

def extract_text(file_path, extension):
    
    text = ''
    if extension == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif extension == '.docx':
        text = extract_text_from_docx(file_path)
    elif extension == '.doc':
        text = extract_text_from_doc(file_path)
    return text



def extract_email(text):

    email = re.findall("([^@|\s]+@[^@]+\.[^@|\s]+)", text)
    if email:
        try:
            return email[0].split()[0].strip(';')
        except IndexError:
            return None

def extract_name(nlp_text, matcher):
  
    pattern = [cs.NAME_PATTERN]

    matcher.add('NAME', None, *pattern)

    matches = matcher(nlp_text)

    for _, start, end in matches:
        span = nlp_text[start:end]
        if 'name' not in span.text.lower():
            return span.text



def extract_mobile_number(text, custom_regex=None):

    if not custom_regex:
        mob_num_regex = r'''(\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)
                        [-\.\s]*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})'''
        phone = re.findall(re.compile(mob_num_regex), text)
    else:
        phone = re.findall(re.compile(custom_regex), text)
    if phone:
        number = ''.join(phone[0])
        return number


def extract_skills(nlp_text, noun_chunks):
    
    tokens = [token.text for token in nlp_text if not token.is_stop]
    data = pd.read_csv(os.path.join(os.path.dirname(__file__), 'skills.csv')) 
    skills = list(data.columns.values)
    skillset = []
    # check for one-grams
    for token in tokens:
        if token.lower() in skills:
            skillset.append(token)
    
    # check for bi-grams and tri-grams
    for token in noun_chunks:
        token = token.text.lower().strip()
        if token in skills:
            skillset.append(token)
    return [i.capitalize() for i in set([i.lower() for i in skillset])]

def cleanup(token, lower = True):
    if lower:
       token = token.lower()
    return token.strip()


def get_total_experience(experience_list):
    
    exp_ = []
    for line in experience_list:
        experience = re.search(
            r'(?P<fmonth>\w+.\d+)\s*(\D|to)\s*(?P<smonth>\w+.\d+|present)',
            line,
            re.I
        )
        if experience:
            exp_.append(experience.groups())
    total_exp = sum(
        [get_number_of_months_from_dates(i[0], i[2]) for i in exp_]
    )
    total_experience_in_months = total_exp
    return total_experience_in_months


def get_number_of_months_from_dates(date1, date2):
    
    if date2.lower() == 'present':
        date2 = datetime.now().strftime('%b %Y')
    try:
        if len(date1.split()[0]) > 3:
            date1 = date1.split()
            date1 = date1[0][:3] + ' ' + date1[1]
        if len(date2.split()[0]) > 3:
            date2 = date2.split()
            date2 = date2[0][:3] + ' ' + date2[1]
    except IndexError:
        return 0
    try:
        date1 = datetime.strptime(str(date1), '%b %Y')
        date2 = datetime.strptime(str(date2), '%b %Y')
        months_of_experience = relativedelta.relativedelta(date2, date1)
        months_of_experience = (months_of_experience.years
                                * 12 + months_of_experience.months)
    except ValueError:
        return 0
    return months_of_experience

def extract_entities_wih_custom_model(custom_nlp_text):

    entities = {}
    for ent in custom_nlp_text.ents:
        if ent.label_ not in entities.keys():
            entities[ent.label_] = [ent.text]
        else:
            entities[ent.label_].append(ent.text)
    for key in entities.keys():
        entities[key] = entities[key]
    return entities

def os_extraction(text):
    resume = text

    postingSkip = random.randint(1,350)

    postingInfo = requests.get(
        f'https://api.lever.co/v0/postings/leverdemo?skip={postingSkip}&limit=1&mode=json'
    )

    postingURL = postingInfo.json()[0]['applyUrl']

    posting = requests.get(postingURL)
    root = html.fromstring(posting.text)
    csrf = root.get_element_by_id("csrf-token").get('value')
    postingID = root.get_element_by_id("posting-id").get('value')

    headers = {
        'Referer': 'https://jobs.lever.co/',
        'Origin': 'https://jobs.lever.co/'
    }

    parseResponse = requests.post(
        'https://jobs.lever.co/parseResume', 
        files=dict(
            resume=('resume.pdf', resume),
            csrf=(None,csrf),
            postingId=(None,postingID)
            ),
        headers=headers
        )

    response = json.dumps(parseResponse.json()) 

    json_obj = json.loads(response)

    return json_obj

def extract_entity_sections_grad(text):
    '''
    Helper function to extract all the raw text from sections of
    resume specifically for graduates and undergraduates
    :param text: Raw text of resume
    :return: dictionary of entities
    '''
    text_split = [i.strip() for i in text.split('\n')]
    # sections_in_resume = [i for i in text_split if i.lower() in sections]
    entities = {}
    key = False
    for phrase in text_split:
        if len(phrase) == 1:
            p_key = phrase
        else:
            p_key = set(phrase.lower().split()) & set(cs.RESUME_SECTIONS_GRAD)
        try:
            p_key = list(p_key)[0]
        except IndexError:
            pass
        if p_key in cs.RESUME_SECTIONS_GRAD:
            entities[p_key] = []
            key = p_key
        elif key and phrase.strip():
            entities[key].append(phrase)

    # entity_key = False
    # for entity in entities.keys():
    #     sub_entities = {}
    #     for entry in entities[entity]:
    #         if u'\u2022' not in entry:
    #             sub_entities[entry] = []
    #             entity_key = entry
    #         elif entity_key:
    #             sub_entities[entity_key].append(entry)
    #     entities[entity] = sub_entities

    # pprint.pprint(entities)

    # make entities that are not found None
    # for entity in cs.RESUME_SECTIONS:
    #     if entity not in entities.keys():
    #         entities[entity] = None
    return entities

def extract_experience(lst):

    durations = []
    company_names = []
    designations = []
    summary = []


    for i in lst:
        company_names.append(i['org'])
        designations.append(i['title'])
        summary.append(i['summary'])
        if 'start' in i:
            start = str(i['start']['month']) + '/' + str(i['start']['year'])
        else:
            start = '?'
        if i['isCurrent'] == True:
            end = 'present'
        else:
            if 'end' in i:
                end = str(i['end']['month']) + '/' + str(i['end']['year'])
            else:
                end = '?'
        
        durations.append(start + ' - ' + end)
    
    exp = list(zip(company_names,designations,durations,summary))
    
    return to_json_string(exp)
    

def to_json_string(lst):

    l = ["company_name","title","duration","summary"]
    d = {}
    for i in range(len(lst)): 
        t = {}
        for j in range(len(l)):
            t[l[j]] = lst[i][j]
        d[i] = t

    return json.dumps(d)


        



